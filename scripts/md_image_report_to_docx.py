#!/usr/bin/env python
from __future__ import annotations

import argparse
import hashlib
import json
import re
import sys
import zipfile
from pathlib import Path
from urllib.parse import unquote


IMAGE_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
ORDERED_RE = re.compile(r"^\s*\d+[.)]\s+(.+?)\s*$")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Convert a Markdown image report into a DOCX with embedded images."
    )
    parser.add_argument("markdown", type=Path, help="Markdown report path.")
    parser.add_argument("--out", type=Path, help="Output DOCX path.")
    parser.add_argument(
        "--image-root",
        type=Path,
        help="Optional root used to resolve relative image paths before the markdown folder.",
    )
    parser.add_argument(
        "--compress",
        action="store_true",
        help="Compress images into a sibling cache folder before embedding.",
    )
    parser.add_argument("--cache-dir", type=Path, help="Optional image compression cache.")
    parser.add_argument("--max-image-px", type=int, default=1400)
    parser.add_argument("--jpeg-quality", type=int, default=82)
    parser.add_argument("--page-image-width", type=float, default=5.8)
    parser.add_argument("--table-image-width", type=float, default=1.35)
    parser.add_argument("--stats-out", type=Path, help="Optional JSON path for conversion stats.")
    parser.add_argument("--fail-on-missing-images", action="store_true", help="Exit with failure when any markdown image cannot be resolved.")
    return parser.parse_args()


def clean_image_target(raw: str) -> str:
    value = raw.strip()
    if value.startswith("<") and ">" in value:
        value = value[1 : value.index(">")]
    else:
        quote_idx = value.find(" \"")
        if quote_idx != -1:
            value = value[:quote_idx]
    return unquote(value.strip())


def resolve_image(raw: str, md_dir: Path, image_root: Path | None) -> Path | None:
    target = clean_image_target(raw)
    if not target:
        return None
    candidate = Path(target)
    if candidate.is_absolute() and candidate.exists():
        return candidate
    candidates = []
    if image_root is not None:
        candidates.append(image_root / target)
    candidates.append(md_dir / target)
    for item in candidates:
        if item.exists():
            return item
    return None


def prepare_image(
    path: Path,
    cache_dir: Path | None,
    compress: bool,
    max_px: int,
    jpeg_quality: int,
) -> Path:
    if not compress:
        return path

    try:
        from PIL import Image, ImageOps
    except ImportError as exc:
        raise RuntimeError("Pillow is required when --compress is used.") from exc

    assert cache_dir is not None
    cache_dir.mkdir(parents=True, exist_ok=True)
    digest = hashlib.sha1(str(path.resolve()).encode("utf-8")).hexdigest()[:12]

    with Image.open(path) as image:
        image = ImageOps.exif_transpose(image)
        has_alpha = image.mode in {"RGBA", "LA"} or (
            image.mode == "P" and "transparency" in image.info
        )
        ext = ".png" if has_alpha else ".jpg"
        out = cache_dir / f"{path.stem}_{digest}{ext}"
        if out.exists():
            return out
        if max(image.size) > max_px:
            image.thumbnail((max_px, max_px))
        if has_alpha:
            image.save(out, "PNG", optimize=True)
        else:
            if image.mode != "RGB":
                image = image.convert("RGB")
            image.save(out, "JPEG", quality=jpeg_quality, optimize=True)
        return out


def split_table_row(line: str) -> list[str]:
    value = line.strip()
    if value.startswith("|"):
        value = value[1:]
    if value.endswith("|"):
        value = value[:-1]
    return [cell.strip() for cell in value.split("|")]


def is_table_separator(line: str) -> bool:
    cells = split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", c.strip()) for c in cells)


def add_inline_content(
    paragraph,
    text: str,
    *,
    md_dir: Path,
    image_root: Path | None,
    cache_dir: Path | None,
    compress: bool,
    max_image_px: int,
    jpeg_quality: int,
    image_width,
    stats: dict,
    table_image_width: float | None = None,
) -> None:
    try:
        from docx.shared import Inches
    except ImportError as exc:
        raise RuntimeError("python-docx is required to create DOCX files.") from exc

    pos = 0
    for match in IMAGE_RE.finditer(text):
        before = text[pos : match.start()]
        if before:
            paragraph.add_run(before)
        stats["markdown_images"] += 1
        resolved = resolve_image(match.group(2), md_dir, image_root)
        if resolved is None:
            stats["missing_images"].append(clean_image_target(match.group(2)))
            paragraph.add_run(f"[missing image: {clean_image_target(match.group(2))}]")
        else:
            prepared = prepare_image(
                resolved, cache_dir, compress, max_image_px, jpeg_quality
            )
            run = paragraph.add_run()
            run.add_picture(str(prepared), width=Inches(image_width))
            stats["embedded_images"] += 1
        pos = match.end()
    tail = text[pos:]
    if tail:
        paragraph.add_run(tail)


def add_markdown_table(doc, lines: list[str], context: dict) -> None:
    rows = []
    for line in lines:
        if is_table_separator(line):
            continue
        rows.append(split_table_row(line))
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    for row_idx, row in enumerate(rows):
        for col_idx in range(col_count):
            cell = table.cell(row_idx, col_idx)
            value = row[col_idx] if col_idx < len(row) else ""
            paragraph = cell.paragraphs[0]
            add_inline_content(
                paragraph,
                value,
                image_width=context["table_image_width"],
                **context,
            )
    doc.add_paragraph()


def is_table_line(line: str) -> bool:
    value = line.strip()
    return value.startswith("|") and value.endswith("|") and "|" in value[1:-1]


def count_docx_parts(path: Path) -> dict:
    with zipfile.ZipFile(path) as archive:
        names = archive.namelist()
        media = [n for n in names if n.startswith("word/media/")]
        try:
            document_xml = archive.read("word/document.xml").decode("utf-8", "ignore")
        except KeyError:
            document_xml = ""
    return {
        "docx_media_files": len(media),
        "docx_drawing_elements": document_xml.count("<w:drawing"),
        "docx_blip_refs": document_xml.count("<a:blip"),
    }


def convert(args: argparse.Namespace) -> dict:
    try:
        from docx import Document
        from docx.shared import Inches, Pt
    except ImportError as exc:
        raise RuntimeError("python-docx is required. Install it with: pip install python-docx") from exc

    md_path = args.markdown.resolve()
    if not md_path.exists():
        raise FileNotFoundError(md_path)
    out_path = (
        args.out.resolve()
        if args.out
        else md_path.with_name(f"{md_path.stem}_embedded_images.docx").resolve()
    )
    out_path.parent.mkdir(parents=True, exist_ok=True)
    cache_dir = (
        args.cache_dir.resolve()
        if args.cache_dir
        else out_path.with_suffix("").with_name(f"{out_path.stem}_image_cache")
    )
    image_root = args.image_root.resolve() if args.image_root else None

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(9)

    stats = {
        "markdown": str(md_path),
        "out": str(out_path),
        "markdown_images": 0,
        "embedded_images": 0,
        "missing_images": [],
    }
    context = {
        "md_dir": md_path.parent,
        "image_root": image_root,
        "cache_dir": cache_dir,
        "compress": args.compress,
        "max_image_px": args.max_image_px,
        "jpeg_quality": args.jpeg_quality,
        "stats": stats,
        "table_image_width": args.table_image_width,
    }

    lines = md_path.read_text(encoding="utf-8").splitlines()
    in_code = False
    idx = 0
    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code = not in_code
            idx += 1
            continue
        if in_code:
            if stripped:
                doc.add_paragraph(line, style="No Spacing")
            idx += 1
            continue
        if not stripped:
            idx += 1
            continue

        if is_table_line(line):
            block = []
            while idx < len(lines) and is_table_line(lines[idx]):
                block.append(lines[idx])
                idx += 1
            add_markdown_table(doc, block, context)
            continue

        heading = HEADING_RE.match(line)
        if heading:
            level = min(len(heading.group(1)), 4)
            doc.add_heading(heading.group(2), level=level)
            idx += 1
            continue

        if stripped in {"---", "***", "___"}:
            doc.add_paragraph("")
            idx += 1
            continue

        bullet = stripped.startswith("- ") or stripped.startswith("* ")
        ordered = ORDERED_RE.match(stripped)
        if bullet:
            paragraph = doc.add_paragraph(style="List Bullet")
            text = stripped[2:].strip()
        elif ordered:
            paragraph = doc.add_paragraph(style="List Number")
            text = ordered.group(1)
        else:
            paragraph = doc.add_paragraph()
            text = line

        add_inline_content(
            paragraph,
            text,
            image_width=args.page_image_width,
            **context,
        )
        idx += 1

    doc.save(out_path)
    stats.update(count_docx_parts(out_path))
    return stats


def main() -> int:
    args = parse_args()
    try:
        stats = convert(args)
    except Exception as exc:
        print(json.dumps({"ok": False, "error": str(exc)}, ensure_ascii=False), file=sys.stderr)
        return 1
    ok = not (args.fail_on_missing_images and stats["missing_images"])
    payload = {"ok": ok, **stats}
    if args.stats_out:
        args.stats_out.resolve().parent.mkdir(parents=True, exist_ok=True)
        args.stats_out.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(payload, ensure_ascii=False, indent=2))
    if not ok:
        return 2
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
